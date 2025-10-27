# chat.py - Hoàn thiện với RAG từ file txt, docx, xlsx, pdf và xử lý multi-sheet
from fastapi import APIRouter, Depends, HTTPException, File, UploadFile, Body
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from app.db import get_db
from app.models import ChatMessage as ModelChatMessage, Conversation as ModelConversation, User
from app.schemas import ChatMessageIn, Conversation, ConversationCreate, ConversationUpdate, ChatMessage, ChatMessageUpdate
from app.routers.task import get_current_user
import ollama
from ollama import web_search, web_fetch
import faiss
import numpy as np
import json
import os
from datetime import datetime
from typing import List, Optional, Dict, Any, Union
import logging
import base64
import PyPDF2
from docx import Document
import pandas as pd
import io
from concurrent.futures import ThreadPoolExecutor
from rank_bm25 import BM25Okapi
import re
import glob

logger = logging.getLogger(__name__)

router = APIRouter()
DIM = 768
executor = ThreadPoolExecutor(max_workers=4)

# Định nghĩa thư mục RAG
RAG_FILES_DIR = "rag_files"
os.makedirs(RAG_FILES_DIR, exist_ok=True)

def get_embedding(text: str, max_length: int = 1024) -> np.ndarray:
    """Tạo embedding cho text, với max_length tăng để giữ nhiều context hơn."""
    try:
        if len(text) > max_length:
            text = text[:max_length]
        resp = ollama.embeddings(model="embeddinggemma:latest", prompt=text)
        return np.array(resp["embedding"])
    except Exception as e:
        logger.error(f"Lỗi khi tạo embedding từ Ollama: {e}")
        return np.zeros(DIM)

def extract_excel_metadata(file_content: bytes) -> Dict[str, Any]:
    """Trích xuất metadata chi tiết từ file Excel"""
    try:
        excel_file = pd.ExcelFile(io.BytesIO(file_content))
        metadata = {
            "file_type": "excel",
            "sheet_count": len(excel_file.sheet_names),
            "sheets": []
        }
        
        for sheet_name in excel_file.sheet_names:
            try:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                sheet_info = {
                    "name": sheet_name,
                    "rows": len(df),
                    "columns": len(df.columns),
                    "columns_list": df.columns.tolist(),
                    "sample_data": df.head(3).to_dict('records') if not df.empty else []
                }
                metadata["sheets"].append(sheet_info)
            except Exception as e:
                logger.warning(f"Failed to analyze sheet {sheet_name}: {e}")
                continue
        
        return metadata
    except Exception as e:
        logger.error(f"Error extracting Excel metadata: {e}")
        return {}

def extract_docx_metadata(file_content: bytes) -> Dict[str, Any]:
    """Trích xuất metadata chi tiết từ file DOCX"""
    try:
        doc = Document(io.BytesIO(file_content))
        metadata = {
            "file_type": "docx",
            "paragraph_count": len([p for p in doc.paragraphs if p.text.strip()]),
            "table_count": len(doc.tables),
            "sections_count": len(doc.sections)
        }
        return metadata
    except Exception as e:
        logger.error(f"Error extracting DOCX metadata: {e}")
        return {}

def extract_text_from_file(file_content: Union[bytes, str]) -> str:
    """Trích xuất nội dung từ file PDF, CSV, DOCX, TXT, XLSX hoặc text."""
    if isinstance(file_content, str):
        try:
            file_content = base64.b64decode(file_content)
        except:
            return file_content

    def sync_extract(file_content: bytes) -> str:
        try:
            # PDF files
            reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = '\n'.join(page.extract_text() for page in reader.pages if page.extract_text())
            if text.strip():
                logger.info("Successfully extracted text from PDF")
                return text[:20000]
        except Exception as e:
            logger.warning(f"PDF extraction failed: {e}")
            pass
        
        try:
            # CSV files
            df = pd.read_csv(io.BytesIO(file_content))
            text = df.to_string(index=False)
            if text.strip():
                logger.info("Successfully extracted text from CSV")
                return text[:20000]
        except Exception as e:
            logger.warning(f"CSV extraction failed: {e}")
            pass
        
        try:
            # Excel files (xlsx, xls) - XỬ LÝ NHIỀU SHEET
            excel_file = pd.ExcelFile(io.BytesIO(file_content))
            all_sheets_text = []
            
            for sheet_name in excel_file.sheet_names:
                try:
                    df = pd.read_excel(excel_file, sheet_name=sheet_name)
                    
                    # Xử lý dữ liệu NaN và chuẩn hóa
                    df = df.fillna('')  # Thay thế NaN bằng chuỗi rỗng
                    
                    # Tạo header cho sheet
                    sheet_header = f"=== SHEET: {sheet_name} ===\n"
                    sheet_header += f"Columns: {', '.join(df.columns.astype(str))}\n"
                    sheet_header += f"Shape: {len(df)} rows x {len(df.columns)} columns\n"
                    sheet_header += "-" * 50 + "\n"
                    
                    # Chuyển dataframe thành text với format đẹp
                    sheet_text = df.to_string(index=False, max_rows=100)  # Giới hạn số dòng
                    
                    all_sheets_text.append(sheet_header + sheet_text)
                    logger.info(f"Extracted sheet: {sheet_name} with {len(df)} rows, {len(df.columns)} columns")
                    
                except Exception as e:
                    logger.warning(f"Failed to extract sheet {sheet_name}: {e}")
                    all_sheets_text.append(f"=== SHEET: {sheet_name} ===\n[Error extracting this sheet: {e}]")
                    continue
            
            if all_sheets_text:
                combined_text = "\n\n".join(all_sheets_text)
                logger.info(f"Successfully extracted {len(all_sheets_text)} sheets from Excel")
                return combined_text[:20000]
                
        except Exception as e:
            logger.warning(f"Excel extraction failed: {e}")
            pass
        
        try:
            # Word documents (docx) - XỬ LÝ TABLE VÀ CÁC ELEMENT KHÁC
            doc = Document(io.BytesIO(file_content))
            all_text = []
            
            # Extract paragraphs
            paragraph_count = 0
            for para in doc.paragraphs:
                if para.text.strip():
                    all_text.append(para.text)
                    paragraph_count += 1
            
            # Extract tables với format rõ ràng
            table_count = 0
            for i, table in enumerate(doc.tables, 1):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_data.append(cell_text)
                    if row_data:
                        table_data.append(" | ".join(row_data))
                
                if table_data:
                    table_header = f"\n--- TABLE {i} ---"
                    all_text.append(table_header)
                    all_text.extend(table_data)
                    table_count += 1
            
            if all_text:
                combined_text = "\n".join(all_text)
                logger.info(f"Successfully extracted {paragraph_count} paragraphs and {table_count} tables from DOCX")
                return combined_text[:20000]
                
        except Exception as e:
            logger.warning(f"DOCX extraction failed: {e}")
            pass
        
        try:
            # Text files (txt)
            text = file_content.decode('utf-8', errors='replace')
            if text.strip():
                logger.info("Successfully extracted text from TXT")
                return text[:20000]
        except Exception as e:
            logger.warning(f"TXT extraction failed: {e}")
            pass
        
        # Fallback: try to decode as text with different encodings
        for encoding in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1', 'utf-16']:
            try:
                text = file_content.decode(encoding, errors='replace')
                if len(text.strip()) > 100:
                    logger.info(f"Successfully decoded text with {encoding}")
                    return text[:20000]
            except:
                continue
        
        logger.warning("Could not extract meaningful text from file")
        return ""

    return executor.submit(sync_extract, file_content).result()

def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
    """Chunk text thành các đoạn nhỏ với overlap."""
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        
        # Tìm điểm cắt tốt (xuống dòng, dấu câu)
        if end < len(text):
            # Ưu tiên cắt tại section boundaries
            section_break = text.rfind('\n===', start, end)
            if section_break != -1 and section_break > start + chunk_size // 2:
                end = section_break
            else:
                # Cắt tại xuống dòng
                line_break = text.rfind('\n', start, end)
                if line_break != -1 and line_break > start + chunk_size // 2:
                    end = line_break + 1
                else:
                    # Cắt tại dấu câu
                    for punctuation in ['. ', '! ', '? ', '。', '！', '？', '; ']:
                        punctuation_pos = text.rfind(punctuation, start, end)
                        if punctuation_pos != -1 and punctuation_pos > start + chunk_size // 2:
                            end = punctuation_pos + len(punctuation)
                            break
        
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        # Di chuyển start với overlap
        start = end - overlap
        if start >= len(text):
            break
    
    return chunks

def process_file_for_rag(file_content: bytes, user_id: int, conversation_id: int, filename: str = "") -> str:
    """Xử lý file để tạo RAG context và lưu vào FAISS index."""
    try:
        # Thêm metadata extraction
        metadata = {}
        file_extension = filename.lower().split('.')[-1] if filename else ""
        
        # Extract metadata dựa trên file type
        if file_extension in ['xlsx', 'xls']:
            metadata = extract_excel_metadata(file_content)
        elif file_extension == 'docx':
            metadata = extract_docx_metadata(file_content)
        
        file_text = extract_text_from_file(file_content)
        if not file_text.strip():
            logger.warning("No text extracted from file for RAG")
            return ""
        
        # Thêm metadata vào context
        metadata_context = ""
        if metadata:
            if metadata.get("file_type") == "excel":
                metadata_context = f"\n[Excel File: {metadata['sheet_count']} sheets"
                for sheet in metadata["sheets"]:
                    metadata_context += f", '{sheet['name']}' ({sheet['rows']} rows, {sheet['columns']} columns)"
                metadata_context += "]"
            elif metadata.get("file_type") == "docx":
                metadata_context = f"\n[DOCX File: {metadata['paragraph_count']} paragraphs, {metadata['table_count']} tables]"
        
        chunks = chunk_text(file_text, chunk_size=600, overlap=75)
        if not chunks:
            logger.warning("No chunks created from file text")
            return ""
        
        index, exists = load_faiss(user_id, conversation_id)
        
        embeddings = []
        valid_chunks = []
        
        for i, chunk in enumerate(chunks):
            if len(chunk.strip()) < 20:  # Tăng ngưỡng để bỏ qua chunks quá ngắn
                continue
                
            # Thêm chunk identifier cho dễ debug
            chunk_with_info = f"[Chunk {i+1}/{len(chunks)}] {chunk}"
            
            emb = get_embedding(chunk_with_info, max_length=512)
            if np.any(emb) and not np.all(emb == 0):  # Kiểm tra embedding hợp lệ
                embeddings.append(emb)
                valid_chunks.append(chunk_with_info)
        
        if embeddings:
            emb_array = np.array(embeddings)
            index.add(emb_array)
            
            faiss_path = get_faiss_path(user_id, conversation_id)
            executor.submit(faiss.write_index, index, faiss_path).result()
            
            logger.info(f"Added {len(embeddings)} chunks from file to RAG index")
            
            result_context = f"{metadata_context}\n[File content loaded: {len(valid_chunks)} chunks, {sum(len(chunk) for chunk in valid_chunks)} characters]"
            if filename:
                result_context = f"[File: {filename}]" + result_context
                
            return result_context
        else:
            logger.warning("No valid embeddings created from file chunks")
            return ""
            
    except Exception as e:
        logger.error(f"Error processing file for RAG: {e}")
        return ""

def load_rag_files_to_conversation(user_id: int, conversation_id: int):
    """Tự động load tất cả file trong thư mục rag_files vào conversation"""
    rag_files = []
    
    supported_patterns = [
        os.path.join(RAG_FILES_DIR, "*.pdf"),
        os.path.join(RAG_FILES_DIR, "*.txt"), 
        os.path.join(RAG_FILES_DIR, "*.docx"),
        os.path.join(RAG_FILES_DIR, "*.xlsx"),
        os.path.join(RAG_FILES_DIR, "*.xls"),
        os.path.join(RAG_FILES_DIR, "*.csv")
    ]
    
    for pattern in supported_patterns:
        rag_files.extend(glob.glob(pattern))
    
    if not rag_files:
        logger.info(f"No RAG files found in {RAG_FILES_DIR}")
        return []
    
    loaded_files = []
    for file_path in rag_files:
        try:
            with open(file_path, "rb") as f:
                file_content = f.read()
            
            filename = os.path.basename(file_path)
            logger.info(f"Processing RAG file: {filename}")
            
            rag_context = process_file_for_rag(file_content, user_id, conversation_id, filename)
            
            if rag_context:
                loaded_files.append({
                    "filename": filename,
                    "path": file_path,
                    "chunks_loaded": rag_context
                })
                logger.info(f"Successfully loaded RAG file: {filename}")
            else:
                logger.warning(f"Failed to load RAG file: {filename}")
                
        except Exception as e:
            logger.error(f"Error loading RAG file {file_path}: {e}")
    
    return loaded_files

def get_faiss_path(user_id: int, conversation_id: int) -> str:
    """Tạo đường dẫn cho FAISS index."""
    index_dir = "faiss_indices"
    os.makedirs(index_dir, exist_ok=True)
    return os.path.join(index_dir, f"faiss_{user_id}_{conversation_id}.index")

def load_faiss(user_id: int, conversation_id: int) -> tuple:
    """Tải hoặc tạo mới FAISS index."""
    path = get_faiss_path(user_id, conversation_id)
    if os.path.exists(path):
        try:
            index = faiss.read_index(path)
            return index, True
        except Exception as e:
            logger.error(f"Error loading FAISS index, creating new: {e}")
    index = faiss.IndexFlatL2(DIM)
    return index, False

def cleanup_faiss_index(user_id: int, conversation_id: int):
    """Dọn dẹp FAISS index nếu tồn tại."""
    path = get_faiss_path(user_id, conversation_id)
    if os.path.exists(path):
        try:
            os.remove(path)
            logger.info(f"Đã xóa FAISS index: {path}")
        except Exception as e:
            logger.error(f"Lỗi khi xóa FAISS index {path}: {e}")

def evaluate_user_input(input_text: str) -> Dict[str, bool]:
    """Đánh giá input của người dùng bằng 4T-Small để xác định model phù hợp."""
    try:
        eval_prompt = f"""
        Bạn là một AI đánh giá input người dùng. Phân tích input sau và trả về JSON với hai trường:
        - "needs_logic": true nếu input liên quan đến toán học, lập trình, hoặc yêu cầu logic phức tạp.
        - "needs_reasoning": true nếu input yêu cầu suy luận sâu, phân tích phức tạp, hoặc trả lời dựa trên lý luận.
        Nếu không thuộc hai trường hợp trên, cả hai trường đều false.
        Input: {input_text}
        Lưu ý: Chỉ trả về định dạng JSON: {{"needs_logic": bool, "needs_reasoning": bool}}, không thêm bất kỳ text nào khác.
        """
        response = ollama.chat(
            model="4T-Small",
            messages=[{"role": "system", "content": eval_prompt}],
            stream=False,
            options={
                "temperature": 0,
                "top_p": 0
            }
        )
        try:
            result = json.loads(response["message"]["content"])
            return {
                "needs_logic": bool(result.get("needs_logic", False)),
                "needs_reasoning": bool(result.get("needs_reasoning", False))
            }
        except json.JSONDecodeError:
            logger.error("Lỗi khi parse JSON từ đánh giá input")
            return {"needs_logic": False, "needs_reasoning": False}
    except Exception as e:
        logger.error(f"Lỗi khi đánh giá input: {e}")
        return {"needs_logic": False, "needs_reasoning": False}

@router.post("/chat", response_class=StreamingResponse)
def chat(
    message: ChatMessageIn = Body(...),
    file: Optional[Union[UploadFile, str]] = Body(None),
    conversation_id: Optional[int] = None,
    user_id: int = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    gender = user.gender
    xung_ho = "anh" if gender == "male" else "chị" if gender == "female" else "bạn"
    logger.debug(f"DEBUG GENDER: {xung_ho}")

    current_time = datetime.now().strftime("%Y-%m-%d %I:%M %p %z")
    
    system_prompt = f"""
      Bạn là Nhi — một AI nói chuyện tự nhiên, thân thiết như người bạn thân của {xung_ho}.  
      Thời điểm hiện tại: {current_time}.

      **Khi cần gọi tool:**
      Trả đúng định dạng JSON, không thêm lời giải thích:
      {{
        "tool_calls": [
          {{
            "type": "function",
            "function": {{
              "name": "web_search",
              "arguments": "{{\\"query\\": \\"optimized query here\\"}}"
            }}
          }}
        ]
      }}

      **Mục tiêu:**
      - Luôn giữ cảm giác tự nhiên, có cảm xúc nhưng không "diễn".
      - Trả lời ngắn, rõ, không lan man.
      - Luôn duy trì cảm giác "người thật nói chuyện" chứ không như máy.
    """

    conversation = None
    is_new_conversation = False
    if conversation_id is not None:
        conversation = db.query(ModelConversation).filter(
            ModelConversation.id == conversation_id,
            ModelConversation.user_id == user_id
        ).first()
        if not conversation:
            raise HTTPException(404, "Conversation not found or not authorized")
    else:
        conversation = ModelConversation(user_id=user_id, created_at=datetime.utcnow())
        db.add(conversation)
        db.flush()
        is_new_conversation = True
    
    # TỰ ĐỘNG LOAD RAG FILES KHI TẠO CONVERSATION MỚI
    if is_new_conversation:
        logger.info(f"New conversation created, loading RAG files...")
        loaded_files = load_rag_files_to_conversation(user_id, conversation.id)
        if loaded_files:
            logger.info(f"Auto-loaded {len(loaded_files)} RAG files to new conversation")

    is_image = False
    file_content = ""
    images = None
    effective_query = message.message
    file_rag_context = ""

    if file:
        file_bytes = None
        filename = ""
        if isinstance(file, UploadFile):
            file_bytes = file.file.read()
            filename = file.filename or ""
            file.file.close()
        else:
            file_bytes = file.encode('utf-8') if not file.startswith('data:') else base64.b64decode(file.split(',')[1]) if ',' in file else base64.b64decode(file)
            filename = "uploaded_file"

        is_image = filename and any(filename.lower().endswith(ext) for ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp'])
        if is_image:
            images = [base64.b64encode(file_bytes).decode('utf-8')]
            effective_query = message.message
        else:
            file_rag_context = process_file_for_rag(file_bytes, user_id, conversation.id, filename)
            file_content = extract_text_from_file(file_bytes)
            effective_query = f"{message.message}\nNội dung file: {file_content}"
            if file_content:
                effective_query += f"\n(File: {filename})"
                
            logger.info(f"File RAG processing completed: {len(file_rag_context) if file_rag_context else 0} context added")

    eval_result = evaluate_user_input(effective_query)
    logger.debug(f"Đánh giá input: {eval_result}")
    
    if is_image:
        model_name = "qwen3-vl:235b-cloud"
        tools = None
    elif eval_result["needs_logic"]:
        model_name = "4T-Logic"
        tools = [web_search, web_fetch]
    elif eval_result["needs_reasoning"]:
        model_name = "4T-Reasoning"
        tools = [web_search, web_fetch]
    else:
        model_name = "4T-Small"
        tools = [web_search, web_fetch]
    logger.debug(f"Model được chọn: {model_name}")

    query_emb = get_embedding(effective_query, max_length=1024)
    history = db.query(ModelChatMessage).filter(
        ModelChatMessage.user_id == user_id,
        ModelChatMessage.conversation_id == conversation.id
    ).order_by(ModelChatMessage.timestamp.asc()).limit(50).all()

    index, exists = load_faiss(user_id, conversation.id)

    valid_history = [h for h in history if h.embedding and json.loads(h.embedding) is not None]
    
    all_context_parts = []
    
    if file_rag_context:
        all_context_parts.append(f"File Context: {file_rag_context}")
    
    if not valid_history:
        context_from_history = ""
    else:
        if not exists or index.ntotal != len(valid_history):
            index = faiss.IndexFlatL2(DIM)
            embs = np.array([json.loads(h.embedding) for h in valid_history])
            if len(embs) > 0:
                index.add(embs)
                executor.submit(faiss.write_index, index, get_faiss_path(user_id, conversation.id)).result()

        index_contents = [h.content for h in valid_history]
        tokenized_contents = [re.findall(r'\w+', content.lower()) for content in index_contents]
        bm25 = BM25Okapi(tokenized_contents)

        query_tokens = re.findall(r'\w+', effective_query.lower())
        bm25_scores = bm25.get_scores(query_tokens)

        if index.ntotal > 0:
            D, I_faiss = index.search(query_emb.reshape(1, -1), k=min(10, index.ntotal))
            faiss_indices = I_faiss[0]
        else:
            faiss_indices = []

        hybrid_scores = {}
        for i, idx in enumerate(faiss_indices):
            if idx < len(bm25_scores):
                hybrid_score = 0.7 * (1 - D[0][i]) + 0.3 * bm25_scores[idx]
                hybrid_scores[idx] = hybrid_score

        reranked_indices = sorted(hybrid_scores, key=hybrid_scores.get, reverse=True)[:5]
        
        context_messages = []
        for idx in reranked_indices:
            if idx < len(valid_history):
                msg = valid_history[idx]
                sim_score = hybrid_scores[idx]
                if sim_score > 0.3:
                    context_messages.append(msg.content)
                    logger.debug(f"Retrieved msg {idx}: score {sim_score:.3f}, content: {msg.content[:50]}...")
        
        context_from_history = "\n".join(context_messages)
        if context_from_history:
            all_context_parts.append(f"History Context:\n{context_from_history}")
    
    final_context = "\n\n".join(all_context_parts)
    if not final_context:
        logger.warning("No relevant context retrieved, using recent history fallback")
        final_context = "\n".join([h.content for h in valid_history[-10:]])

    full_prompt = f"Context: {final_context}\nUser: {effective_query}" if not is_image else effective_query

    def generate_stream():
      yield f"data: {json.dumps({'conversation_id': conversation.id})}\n\n"
      full_response = []
      messages: List[Dict[str, Any]] = [
          {"role": "system", "content": system_prompt},
          {"role": "user", "content": full_prompt}
      ]
      if is_image:
          messages[-1]["images"] = images
      try:
          api_key = os.getenv('OLLAMA_API_KEY')
          if not api_key:
              raise ValueError("OLLAMA_API_KEY env var not set")
          os.environ['OLLAMA_API_KEY'] = api_key

          while True:
              current_message: Dict[str, Any] = {"role": "assistant", "content": ""}
              tool_calls: List[Dict[str, Any]] = []
              stream = ollama.chat(
                  model=model_name,
                  messages=messages,
                  tools=tools,
                  stream=True,
                  options={
                      "temperature": 0.6,
                      "repeat_penalty": 1.2,
                      "num_predict": -1,
                  }
              )
              for chunk in stream:
                  if "message" in chunk:
                      msg_chunk = chunk["message"]
                      if "tool_calls" in msg_chunk and msg_chunk["tool_calls"]:
                          serialized_tool_calls = [
                              {
                                  "function": {
                                      "name": tc["function"]["name"],
                                      "arguments": tc["function"]["arguments"]
                                  }
                              } for tc in msg_chunk["tool_calls"]
                          ]
                          yield f"data: {json.dumps({'tool_calls': serialized_tool_calls})}\n\n"
                          logger.debug(f"Yielded tool_calls: {serialized_tool_calls}")
                          for tc in msg_chunk["tool_calls"]:
                              if "function" in tc:
                                  tool_calls.append(tc)
                      if "content" in msg_chunk and msg_chunk["content"]:
                          delta = msg_chunk["content"].encode('utf-8').decode('utf-8', errors='replace')
                          current_message["content"] += delta
                          full_response.append(delta)
                          yield f"data: {json.dumps({'content': delta})}\n\n"
              messages.append(current_message)
              if tool_calls:
                  for tool_call in tool_calls:
                      function_name = tool_call['function']['name']
                      args = tool_call['function']['arguments']
                      if function_name == 'web_search':
                          result = executor.submit(web_search, **args).result()
                      elif function_name == 'web_fetch':
                          result = executor.submit(web_fetch, **args).result()
                      else:
                          result = f"Tool {function_name} not found"
                      tool_msg = {
                          'role': 'tool',
                          'content': str(result)[:10000],
                          'tool_name': function_name
                      }
                      messages.append(tool_msg)
                      logger.debug(f"Tool {function_name} result: {str(result)[:200]}...")
              else:
                  break
          yield f"data: {json.dumps({'done': True})}\n\n"
          executor.submit(save_after_stream, ''.join(full_response)).result()
      except Exception as e:
          logger.error(f"Lỗi trong stream generation: {e}")
          yield f"data: {json.dumps({'error': str(e)})}\n\n"

    def save_after_stream(full_response: str):
        if not full_response:
            logger.error("Empty response from stream")
            return
        try:
            ass_emb = get_embedding(full_response, max_length=1024)
            user_msg_content = effective_query if not is_image else f"{message.message} (Image: {filename})"
            user_msg = ModelChatMessage(
                user_id=user_id,
                conversation_id=conversation.id,
                content=user_msg_content,
                role="user",
                embedding=json.dumps(query_emb.tolist())
            )
            ass_msg = ModelChatMessage(
                user_id=user_id,
                conversation_id=conversation.id,
                content=full_response,
                role="assistant",
                embedding=json.dumps(ass_emb.tolist())
            )
            db.add_all([user_msg, ass_msg])
            db.flush()
            
            index = faiss.IndexFlatL2(DIM)
            all_msgs = db.query(ModelChatMessage).filter(ModelChatMessage.conversation_id == conversation.id).all()
            valid_embs = [json.loads(m.embedding) for m in all_msgs if m.embedding]
            if valid_embs:
                index.add(np.array(valid_embs))
            executor.submit(faiss.write_index, index, get_faiss_path(user_id, conversation.id)).result()
            db.commit()
        except Exception as e:
            db.rollback()
            logger.error(f"Lỗi khi lưu tin nhắn hoặc cập nhật FAISS index: {e}")

    return StreamingResponse(generate_stream(), media_type="text/event-stream")

@router.post("/rag/load-to-conversation")
async def load_rag_files_to_current_conversation(
    conversation_id: int,
    user_id: int = Depends(get_current_user), 
    db: Session = Depends(get_db)
):
    """Load tất cả file trong thư mục rag_files vào conversation hiện tại"""
    conversation = db.query(ModelConversation).filter(
        ModelConversation.id == conversation_id,
        ModelConversation.user_id == user_id
    ).first()
    
    if not conversation:
        raise HTTPException(404, "Conversation not found")
    
    loaded_files = load_rag_files_to_conversation(user_id, conversation_id)
    
    return {
        "message": f"Loaded {len(loaded_files)} RAG files into conversation",
        "loaded_files": loaded_files
    }

@router.get("/rag/files")
async def list_rag_files():
    """Liệt kê tất cả file có trong thư mục RAG"""
    rag_files = []
    
    supported_extensions = ['.pdf', '.txt', '.docx', '.xlsx', '.xls', '.csv']
    
    for filename in os.listdir(RAG_FILES_DIR):
        file_path = os.path.join(RAG_FILES_DIR, filename)
        if os.path.isfile(file_path) and any(filename.lower().endswith(ext) for ext in supported_extensions):
            file_stats = os.stat(file_path)
            rag_files.append({
                "filename": filename,
                "size": file_stats.st_size,
                "modified_time": datetime.fromtimestamp(file_stats.st_mtime).isoformat()
            })
    
    return {"rag_files": rag_files}

@router.post("/rag/analyze-file")
async def analyze_rag_file(
    file: UploadFile = File(...),
    user_id: int = Depends(get_current_user)
):
    """Phân tích metadata của file RAG"""
    try:
        file_content = await file.read()
        filename = file.filename
        
        metadata = {}
        if filename.lower().endswith(('.xlsx', '.xls')):
            metadata = extract_excel_metadata(file_content)
        elif filename.lower().endswith('.docx'):
            metadata = extract_docx_metadata(file_content)
        
        # Extract sample text
        sample_text = extract_text_from_file(file_content)
        sample_preview = sample_text[:500] + "..." if len(sample_text) > 500 else sample_text
        
        return {
            "filename": filename,
            "metadata": metadata,
            "sample_preview": sample_preview,
            "total_length": len(sample_text)
        }
        
    except Exception as e:
        logger.error(f"Error analyzing file: {e}")
        raise HTTPException(500, f"Error analyzing file: {str(e)}")

@router.post("/conversations", response_model=Conversation)
def create_conversation(conv: ConversationCreate, user_id: int = Depends(get_current_user), db: Session = Depends(get_db)):
    conversation = ModelConversation(user_id=user_id, created_at=datetime.utcnow())
    db.add(conversation)
    db.commit()
    db.refresh(conversation)
    return conversation

@router.get("/conversations", response_model=List[Conversation])
def get_conversations(user_id: int = Depends(get_current_user), db: Session = Depends(get_db)):
    conversations = db.query(ModelConversation).filter(ModelConversation.user_id == user_id).all()
    return conversations

@router.get("/conversations/{id}", response_model=Conversation)
def get_conversation(id: int, user_id: int = Depends(get_current_user), db: Session = Depends(get_db)):
    conversation = db.query(ModelConversation).filter(ModelConversation.id == id, ModelConversation.user_id == user_id).first()
    if not conversation:
        raise HTTPException(404, "Conversation not found or not authorized")
    return conversation

@router.put("/conversations/{id}", response_model=Conversation)
def update_conversation(id: int, conv_update: ConversationUpdate, user_id: int = Depends(get_current_user), db: Session = Depends(get_db)):
    conversation = db.query(ModelConversation).filter(ModelConversation.id == id, ModelConversation.user_id == user_id).first()
    if not conversation:
        raise HTTPException(404, "Conversation not found or not authorized")
    db.commit()
    db.refresh(conversation)
    return conversation

@router.delete("/conversations/{id}")
def delete_conversation(id: int, user_id: int = Depends(get_current_user), db: Session = Depends(get_db)):
    conversation = db.query(ModelConversation).filter(ModelConversation.id == id, ModelConversation.user_id == user_id).first()
    if not conversation:
        raise HTTPException(404, "Conversation not found or not authorized")

    index_path = get_faiss_path(user_id, id)
    if os.path.exists(index_path):
        os.remove(index_path)

    db.delete(conversation)
    db.commit()
    return {"message": "Conversation deleted"}

@router.get("/conversations/{conversation_id}/messages", response_model=List[ChatMessage])
def get_messages(conversation_id: int, user_id: int = Depends(get_current_user), db: Session = Depends(get_db)):
    conversation = db.query(ModelConversation).filter(ModelConversation.id == conversation_id, ModelConversation.user_id == user_id).first()
    if not conversation:
        raise HTTPException(404, "Conversation not found or not authorized")

    messages = db.query(ModelChatMessage).filter(
        ModelChatMessage.conversation_id == conversation_id,
        ModelChatMessage.user_id == user_id
    ).order_by(ModelChatMessage.timestamp.asc()).all()

    result = []
    for msg in messages:
        msg_dict = msg.__dict__
        if msg.embedding and isinstance(msg.embedding, str):
            try:
                parsed_embedding = json.loads(msg.embedding)
                msg_dict['embedding'] = parsed_embedding
            except json.JSONDecodeError:
                msg_dict['embedding'] = None
        result.append(ChatMessage(**msg_dict))

    return result

@router.get("/messages/{id}", response_model=ChatMessage)
def get_message(id: int, user_id: int = Depends(get_current_user), db: Session = Depends(get_db)):
    message = db.query(ModelChatMessage).filter(ModelChatMessage.id == id, ModelChatMessage.user_id == user_id).first()
    if not message:
        raise HTTPException(404, "Message not found or not authorized")

    msg_dict = message.__dict__
    if msg_dict['embedding'] and isinstance(msg_dict['embedding'], str):
        try:
            parsed_embedding = json.loads(msg_dict['embedding'])
            msg_dict['embedding'] = parsed_embedding
        except json.JSONDecodeError:
            msg_dict['embedding'] = None
    return ChatMessage(**msg_dict)

@router.put("/messages/{id}", response_model=ChatMessage)
def update_message(id: int, msg_update: ChatMessageUpdate, user_id: int = Depends(get_current_user), db: Session = Depends(get_db)):
    message = db.query(ModelChatMessage).filter(ModelChatMessage.id == id, ModelChatMessage.user_id == user_id).first()
    if not message:
        raise HTTPException(404, "Message not found or not authorized")

    if msg_update.content is not None:
        message.content = msg_update.content
        new_embedding = get_embedding(msg_update.content, max_length=1024)
        message.embedding = json.dumps(new_embedding.tolist())

        index, _ = load_faiss(user_id, message.conversation_id)
        all_messages = db.query(ModelChatMessage).filter(ModelChatMessage.conversation_id == message.conversation_id).all()
        embs = np.array([json.loads(m.embedding) for m in all_messages if m.embedding])

        index.reset()
        if len(embs) > 0:
            index.add(embs)
        faiss.write_index(index, get_faiss_path(user_id, message.conversation_id))

    db.commit()
    db.refresh(message)
    msg_dict = message.__dict__
    if msg_dict['embedding'] and isinstance(msg_dict['embedding'], str):
        try:
            parsed_embedding = json.loads(msg_dict['embedding'])
            msg_dict['embedding'] = parsed_embedding
        except json.JSONDecodeError:
            msg_dict['embedding'] = None
    return ChatMessage(**msg_dict)

@router.delete("/messages/{id}")
def delete_message(id: int, user_id: int = Depends(get_current_user), db: Session = Depends(get_db)):
    message = db.query(ModelChatMessage).filter(ModelChatMessage.id == id, ModelChatMessage.user_id == user_id).first()
    if not message:
        raise HTTPException(404, "Message not found or not authorized")

    db.delete(message)
    db.commit()

    index, _ = load_faiss(user_id, message.conversation_id)
    embs = np.array([json.loads(m.embedding) for m in db.query(ModelChatMessage).filter(ModelChatMessage.conversation_id == message.conversation_id).all() if m.embedding])

    index.reset()
    if len(embs) > 0:
        index.add(embs)
    faiss.write_index(index, get_faiss_path(user_id, message.conversation_id))

    return {"message": "Message deleted"}
