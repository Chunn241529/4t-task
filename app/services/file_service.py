import io
import pandas as pd
import PyPDF2
from docx import Document
import base64
import logging
from typing import Dict, Any, Union
from concurrent.futures import ThreadPoolExecutor

logger = logging.getLogger(__name__)

class FileService:
    def __init__(self):
        self.executor = ThreadPoolExecutor(max_workers=4)

    def extract_excel_metadata(self, file_content: bytes) -> Dict[str, Any]:
        """Trích xuất metadata chi tiết từ file Excel"""
        try:
            excel_file = pd.ExcelFile(io.BytesIO(file_content))
            metadata = {
                "file_type": "excel",
                "sheet_count": len(excel_file.sheet_names),
                "sheets": []
            }
            
            for sheet_name in excel_file.sheet_names:
                try:
                    df = pd.read_excel(excel_file, sheet_name=sheet_name)
                    sheet_info = {
                        "name": sheet_name,
                        "rows": len(df),
                        "columns": len(df.columns),
                        "columns_list": df.columns.tolist(),
                        "sample_data": df.head(3).to_dict('records') if not df.empty else []
                    }
                    metadata["sheets"].append(sheet_info)
                except Exception as e:
                    logger.warning(f"Failed to analyze sheet {sheet_name}: {e}")
                    continue
            
            return metadata
        except Exception as e:
            logger.error(f"Error extracting Excel metadata: {e}")
            return {}

    def extract_docx_metadata(self, file_content: bytes) -> Dict[str, Any]:
        """Trích xuất metadata chi tiết từ file DOCX"""
        try:
            doc = Document(io.BytesIO(file_content))
            metadata = {
                "file_type": "docx",
                "paragraph_count": len([p for p in doc.paragraphs if p.text.strip()]),
                "table_count": len(doc.tables),
                "sections_count": len(doc.sections)
            }
            return metadata
        except Exception as e:
            logger.error(f"Error extracting DOCX metadata: {e}")
            return {}

    def extract_text_from_file(self, file_content: Union[bytes, str]) -> str:
        """Trích xuất nội dung từ file PDF, CSV, DOCX, TXT, XLSX hoặc text."""
        if isinstance(file_content, str):
            try:
                file_content = base64.b64decode(file_content)
            except:
                return file_content

        def sync_extract(file_content: bytes) -> str:
            try:
                # PDF files
                reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                text = '\n'.join(page.extract_text() for page in reader.pages if page.extract_text())
                if text.strip():
                    logger.info("Successfully extracted text from PDF")
                    return text[:20000]
            except Exception as e:
                logger.warning(f"PDF extraction failed: {e}")
                pass
            
            try:
                # CSV files
                df = pd.read_csv(io.BytesIO(file_content))
                text = df.to_string(index=False)
                if text.strip():
                    logger.info("Successfully extracted text from CSV")
                    return text[:20000]
            except Exception as e:
                logger.warning(f"CSV extraction failed: {e}")
                pass
            
            try:
                # Excel files (xlsx, xls) - XỬ LÝ NHIỀU SHEET
                excel_file = pd.ExcelFile(io.BytesIO(file_content))
                all_sheets_text = []
                
                for sheet_name in excel_file.sheet_names:
                    try:
                        df = pd.read_excel(excel_file, sheet_name=sheet_name)
                        df = df.fillna('')
                        
                        sheet_header = f"=== SHEET: {sheet_name} ===\n"
                        sheet_header += f"Columns: {', '.join(df.columns.astype(str))}\n"
                        sheet_header += f"Shape: {len(df)} rows x {len(df.columns)} columns\n"
                        sheet_header += "-" * 50 + "\n"
                        
                        sheet_text = df.to_string(index=False, max_rows=100)
                        all_sheets_text.append(sheet_header + sheet_text)
                        logger.info(f"Extracted sheet: {sheet_name} with {len(df)} rows, {len(df.columns)} columns")
                        
                    except Exception as e:
                        logger.warning(f"Failed to extract sheet {sheet_name}: {e}")
                        all_sheets_text.append(f"=== SHEET: {sheet_name} ===\n[Error extracting this sheet: {e}]")
                        continue
                
                if all_sheets_text:
                    combined_text = "\n\n".join(all_sheets_text)
                    logger.info(f"Successfully extracted {len(all_sheets_text)} sheets from Excel")
                    return combined_text[:20000]
                    
            except Exception as e:
                logger.warning(f"Excel extraction failed: {e}")
                pass
            
            try:
                # Word documents (docx)
                doc = Document(io.BytesIO(file_content))
                all_text = []
                
                paragraph_count = 0
                for para in doc.paragraphs:
                    if para.text.strip():
                        all_text.append(para.text)
                        paragraph_count += 1
                
                table_count = 0
                for i, table in enumerate(doc.tables, 1):
                    table_data = []
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_data.append(cell_text)
                        if row_data:
                            table_data.append(" | ".join(row_data))
                    
                    if table_data:
                        table_header = f"\n--- TABLE {i} ---"
                        all_text.append(table_header)
                        all_text.extend(table_data)
                        table_count += 1
                
                if all_text:
                    combined_text = "\n".join(all_text)
                    logger.info(f"Successfully extracted {paragraph_count} paragraphs and {table_count} tables from DOCX")
                    return combined_text[:20000]
                    
            except Exception as e:
                logger.warning(f"DOCX extraction failed: {e}")
                pass
            
            try:
                # Text files (txt)
                text = file_content.decode('utf-8', errors='replace')
                if text.strip():
                    logger.info("Successfully extracted text from TXT")
                    return text[:20000]
            except Exception as e:
                logger.warning(f"TXT extraction failed: {e}")
                pass
            
            for encoding in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1', 'utf-16']:
                try:
                    text = file_content.decode(encoding, errors='replace')
                    if len(text.strip()) > 100:
                        logger.info(f"Successfully decoded text with {encoding}")
                        return text[:20000]
                except:
                    continue
            
            logger.warning("Could not extract meaningful text from file")
            return ""

        return self.executor.submit(sync_extract, file_content).result()
